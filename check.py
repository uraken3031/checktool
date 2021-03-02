#Part6
import pyperclip
import docx
import glob
import re
from collections import deque
import os, time
import random

#ファイルの中身を取得する
def get_text(fname):
    doc = docx.Document(fname)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

#既存の記事に対して、リライトしたテキストをドキュメントに挿入
def insert_text():
    for fname in glob.glob('lafool*.docx'):
        q = get_text(fname)
        res = q.replace('一つ目','1つ目').replace('繋がる','つながる').replace('たとえば','例えば').replace('わかる','分かる').replace('〜の一つ','〜のひとつ').replace('取組み','取り組み')
        newdoc = docx.Document(fname)
        newdoc.add_paragraph(res)
        print(newdoc.add_paragraph(res))
        newdoc.save(fname)
